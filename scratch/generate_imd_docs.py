from __future__ import annotations

import csv
import json
import re
import textwrap
from collections import Counter
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Image as RLImage
from reportlab.platypus import PageBreak, Paragraph, Preformatted, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[1]
BACKEND = ROOT / "backend"
FRONTEND = ROOT / "frontend"
SRC = FRONTEND / "src"
TODAY = date.today().isoformat()
SCREENSHOT_DIR = ROOT / "docs" / "screenshots"
SELECTED_SCREENSHOT_LIMIT = 14

SCREENSHOT_CAPTIONS = {
    "imd_evidence_01.png": ("Historical Workbench", "Historical thunderstorm analysis center with file/date/station selectors and IMD review readiness panel."),
    "imd_evidence_02.png": ("Historical Workbench", "Review readiness and verification status panel with contingency metrics and analog confidence."),
    "imd_evidence_03.png": ("Historical Workbench", "Pipeline evidence trail showing data ingestion and sounding parsing sections."),
    "imd_evidence_04.png": ("Historical Workbench", "Index calculation cards for CAPE, CIN, LI, PWAT, SWEAT, K-index, shear, and theta-e."),
    "imd_evidence_05.png": ("Historical Workbench", "Collapsed workflow stages for threshold comparison, verification, interpretation, and recommendation."),
    "imd_evidence_06.png": ("Forecast Simulator", "Custom convective sounding forecast lab with thermodynamic slider inputs and simulated forecast outcome."),
    "imd_evidence_07.png": ("Forecast Verification", "Convective index threshold verification lab with threshold controls and contingency metrics."),
    "imd_evidence_08.png": ("AI Prediction Engine", "Probabilistic forecast envelope and operational analog match evidence."),
    "imd_evidence_09.png": ("AI Prediction Engine", "Operational meteorologist decision support and composite convective severity score."),
    "imd_evidence_10.png": ("Climatology & Research Center", "Climatology and skill intelligence station view with seasonal recurrence composites."),
    "imd_evidence_11.png": ("Climatology & Research Center", "Expanded climatology view showing station composite metrics and chart area."),
    "imd_evidence_12.png": ("Coastal Monitoring", "Bay of Bengal coastal radar map focused on north coastal Andhra and station comparison."),
    "imd_evidence_13.png": ("Coastal Monitoring", "Coastal Andhra intelligence stream with marine inflow and corridor metrics."),
    "imd_evidence_14.png": ("Research Navigation", "Research & Soundings navigation stack with Dataset Explorer and Reviewer Audit Dashboard."),
    "imd_evidence_15.png": ("Dataset Explorer", "IMD Convective Dataset Explorer distribution and ingestion audit panels."),
    "imd_evidence_16.png": ("Reviewer Dashboard", "Reviewer audit dashboard empty-state prompt for historical analysis activation."),
    "imd_evidence_17.png": ("Forecast Verification", "Verification lab loading state while contingency matrices are calculated."),
    "imd_evidence_18.png": ("Forecast Verification", "Verification lab controls and calculation state in the Research & Soundings shell."),
    "imd_evidence_19.png": ("Dataset Explorer", "Dataset Explorer metric cards and seeded historical records log."),
    "imd_evidence_20.png": ("Reviewer Dashboard", "Reviewer audit dashboard prompt and reviewer workflow container."),
    "imd_evidence_21.png": ("Forecast Verification", "Verification lab with computed CSI, POD, FAR, BIAS, HSS and case inspection log."),
    "imd_evidence_22.png": ("Forecast Verification", "Computed verification lab with contingency matrix and operational decision advice."),
    "imd_evidence_23.png": ("File Analysis Center", "Uploaded dataset quality summary showing records, missing values, duplicates, quality score, and station coverage."),
    "imd_evidence_24.png": ("Live Doppler Radar", "IMD Visakhapatnam Cyclone Warning Centre radar console with operational metadata HUD."),
    "imd_evidence_25.png": ("AI Prediction Engine", "Regional threat queue, CAPE traceability, probability evolution, and recommended actions."),
    "imd_evidence_26.png": ("AI Prediction Engine", "Advanced diagnostics and scientific calibration blocks behind the operational summary."),
    "imd_evidence_27.png": ("AI Prediction Engine", "Composite severity, lifecycle timeline, and decision-weight calibration controls."),
    "imd_evidence_28.png": ("Convective Analytics", "Convective verification and threshold discovery hub with operational validation curves."),
    "imd_evidence_29.png": ("Convective Analytics", "Historical convective sounding log table from CWC archives."),
    "imd_evidence_30.png": ("Operational Event Archive", "Operational event archive cards for replaying historical severe weather scenarios."),
    "imd_evidence_31.png": ("Skew-T Sounding Plot", "Thermodynamic Skew-T Log-P sounding profile with metadata and station selection."),
    "imd_evidence_32.png": ("Wyoming Sounding Data", "University of Wyoming sounding archive view with source, cycle, cache age, and parameter decoder."),
    "imd_evidence_33.png": ("Historical Thunderstorm Archive", "Archive summary with latest observed thunderstorm event, record counts, and latest record metadata."),
    "imd_evidence_34.png": ("Forecast Simulator", "Custom sounding simulator with CAPE static-data warning and threshold comparison."),
    "imd_evidence_35.png": ("Dataset Explorer", "CAPE dynamicity audit table and latest archive metrics."),
    "imd_evidence_36.png": ("Reviewer Dashboard", "Reviewer audit dashboard with docket exports and reviewer verdict form."),
    "imd_evidence_37.png": ("Research Navigation", "Research & Soundings navigation with Start Here, About, archive, simulator, verification, and learning pages."),
    "imd_evidence_38.png": ("Atmospheric Intelligence Center", "Operational atmospheric intelligence center with station risk cards and interpretation."),
    "imd_evidence_39.png": ("Atmospheric Intelligence Center", "Compact atmospheric intelligence view answering current situation and recommended context."),
    "imd_evidence_40.png": ("Severe Weather Watchdesk", "Severe weather watchdesk with thunderstorm, heavy rain, lightning, squall, and stabilization indicators."),
    "imd_evidence_41.png": ("Thermodynamic Lab", "Thermodynamic lab with CAPE, CIN, LCL, LFC, EL, parcel trace narrative, interpretation, and actions."),
    "imd_evidence_42.png": ("Climatology & Research Center", "Climatology and research center with historical trust, Brier score, interpretation, and actions."),
    "imd_evidence_43.png": ("Thunderstorm Forecast Simulator", "Operational forecast workspace with file analysis center steps and forecast metric cards."),
    "imd_evidence_44.png": ("Live Operational Nowcast Center", "Nowcast center with coastal map markers and probability bars for thunderstorm, lightning, rain, and squall."),
    "imd_evidence_45.png": ("Historical Thunderstorm Archive", "Historical archive summary with file analysis center and latest thunderstorm evidence."),
    "imd_evidence_46.png": ("Forecast Simulator", "Simulator validation stages and CAPE static data warning evidence."),
}


def read(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def count_csv(path: Path) -> int:
    if not path.exists():
        return 0
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        return sum(1 for _ in csv.DictReader(handle))


def parse_routes() -> list[dict]:
    text = read(BACKEND / "main.py").splitlines()
    routes = []
    for idx, line in enumerate(text):
        match = re.search(r'@app\.(get|post|put|delete|websocket)\("([^"]+)"', line)
        if not match:
            continue
        method, route = match.group(1).upper(), match.group(2)
        func = ""
        for inner in text[idx + 1: idx + 12]:
            func_match = re.search(r"def\s+([A-Za-z_][A-Za-z0-9_]*)\(", inner)
            async_match = re.search(r"async\s+def\s+([A-Za-z_][A-Za-z0-9_]*)\(", inner)
            found = func_match or async_match
            if found:
                func = found.group(1)
                break
        routes.append({"method": "WEBSOCKET" if method == "WEBSOCKET" else method, "path": route, "handler": func, "line": idx + 1})
    return routes


def route_purpose(route: dict) -> str:
    path = route["path"]
    handler = route["handler"].replace("_", " ")
    known = {
        "/auth/signup": "Register an operational user in the StormSense authentication store.",
        "/auth/login": "Authenticate a reviewer/operator and return a bearer token plus user profile.",
        "/auth/me": "Return the authenticated operator profile for session validation.",
        "/forecast": "Generate live station forecasts using sounding, thermodynamic, probability, and decision-support engines.",
        "/history": "Return recent persisted thunderstorm forecast records.",
        "/trend-analysis": "Return station-level trend analysis for convective evolution monitoring.",
        "/storm-escalation": "Return severe escalation records for high-risk station monitoring.",
        "/stream/atmospheric": "Stream atmospheric forecast, trend, escalation, and cycle metadata over websocket.",
        "/system-status": "Return backend, database, websocket, and runtime health status.",
        "/cwc/analyze-historical-dataset": "Analyze uploaded CSV/XLS/XLSX historical datasets and build file-based registry output.",
        "/cwc/latest-file-analyzed": "Return the latest uploaded file analysis metadata for dashboard reference.",
        "/cwc/cape-traceability": "Return station CAPE traceability timeline and static-data warnings.",
        "/cwc/historical-dates": "Return the historical archive index used by search, registry, and reviewer workflows.",
        "/cwc/historical-analysis": "Return a full historical case analysis with forecast reproduction and verification fields.",
        "/cwc/district-impact": "Return district impact intelligence for coastal Andhra monitoring.",
        "/cwc/operational-bulletins": "Generate operational bulletin products from current forecast rows.",
    }
    return known.get(path, f"Expose {handler} workflow output for the operational workstation.")


def parse_schema() -> dict[str, list[str]]:
    schema = read(ROOT / "supabase_schema.sql")
    tables: dict[str, list[str]] = {}
    for match in re.finditer(r"CREATE TABLE IF NOT EXISTS\s+([A-Za-z_][A-Za-z0-9_]*)\s*\((.*?)\);", schema, re.S):
        name = match.group(1)
        body = match.group(2)
        cols = []
        for raw in body.splitlines():
            line = raw.strip().rstrip(",")
            if not line or line.startswith("--"):
                continue
            token = line.split()[0]
            if token.upper() in {"PRIMARY", "FOREIGN", "CONSTRAINT", "UNIQUE", "CHECK"}:
                continue
            cols.append(line)
        tables[name] = cols
    return tables


def parse_sidebar() -> list[dict]:
    text = read(SRC / "components" / "layout" / "Sidebar.jsx")
    items = []
    for match in re.finditer(r'id:\s*"([^"]+)",\s*label:\s*"([^"]+)"', text):
        items.append({"id": match.group(1), "label": match.group(2)})
    return items


def parse_module_titles() -> dict[str, str]:
    text = read(SRC / "components" / "modules" / "Phase3OpsModule.jsx")
    block = re.search(r"const MODULE_TITLES = \{(.*?)\};", text, re.S)
    titles = {}
    if block:
        for match in re.finditer(r'([A-Z_]+):\s*"([^"]+)"', block.group(1)):
            titles[match.group(1)] = match.group(2)
    return titles


def project_counts() -> dict:
    files = []
    for base in [BACKEND, SRC]:
        for path in base.rglob("*"):
            if path.is_file() and "venv" not in path.parts and "node_modules" not in path.parts and "dist" not in path.parts:
                files.append(path)
    extensions = Counter(path.suffix or "(none)" for path in files)
    return {
        "total_files": len(files),
        "extensions": dict(sorted(extensions.items())),
        "historical_archive_records": count_csv(BACKEND / "data" / "rsrw_historical_archive.csv"),
        "observational_records": count_csv(BACKEND / "data" / "imd_observational_records.csv"),
        "image_assets": [str(Path("frontend/src/assets/hero.png"))] if (SRC / "assets" / "hero.png").exists() else [],
    }


ROUTES = parse_routes()
SCHEMA = parse_schema()
SIDEBAR = parse_sidebar()
MODULE_TITLES = parse_module_titles()
COUNTS = project_counts()


def screenshot_inventory() -> list[dict]:
    items = []
    if not SCREENSHOT_DIR.exists():
        return items
    for path in sorted(SCREENSHOT_DIR.glob("imd_evidence_*.png")):
        module, caption = SCREENSHOT_CAPTIONS.get(path.name, ("StormSense AI UI", "Supplied UI screenshot evidence."))
        items.append({
            "id": path.stem.replace("imd_evidence_", "IMG-"),
            "file": path,
            "module": module,
            "caption": caption,
            "relative": path.relative_to(ROOT).as_posix(),
        })
    return items


SCREENSHOTS = screenshot_inventory()
COUNTS["image_assets"] = COUNTS["image_assets"] + [item["relative"] for item in SCREENSHOTS]

RESEARCH_TABS = [
    ("START_HERE", "START HERE", "Reviewer onboarding center for data requirements, outputs, and common review flow."),
    ("ABOUT", "ABOUT STORMSENSE AI", "Operational reference page describing purpose, architecture, data sources, and workflows."),
    ("SKEW_T", "Skew-T / Thermodynamic Lab", "Sounding visualization and thermodynamic parameter interpretation."),
    ("WYOMING_DATA", "Wyoming Data", "Radiosonde ingest and cycle-aware sounding review."),
    ("HISTORICAL_WORKBENCH", "Historical Thunderstorm Archive", "Flagship investigation console for 2023, 2024, and 2025 archive records."),
    ("FORECAST_LAB", "Thunderstorm Forecast Simulator", "Historical and custom sounding forecast reproduction and dataset analysis."),
    ("RESEARCH_VERIFY", "Forecast Verification", "Contingency metrics, threshold testing, and reviewer verification workflows."),
    ("DATASET_EXPLORER", "Historical Weather Database", "Historical weather database and archive exploration workspace."),
    ("REVIEWER_DASHBOARD", "IMD Review Dashboard", "Reviewer docket, verdict, and operational review workspace."),
    ("INDEX_GLOSSARY", "Thunderstorm Indices Guide", "Guide to CAPE, LI, SWEAT, PWAT, K-index, CIN, LCL, LFC, EL, and verification metrics."),
    ("TERMINOLOGY", "Meteorology Learning Center", "Operational terminology and learning support for reviewers."),
]

GLOSSARY = [
    ("CAPE", "Convective Available Potential Energy. High CAPE indicates buoyant energy available for deep convection. Operationally, rising CAPE with sufficient moisture and weak inhibition increases thunderstorm potential."),
    ("CIN", "Convective Inhibition. CIN represents the cap suppressing parcel ascent. Strong CIN can delay initiation even with high CAPE; weakening CIN can permit rapid thunderstorm development."),
    ("LI", "Lifted Index. Negative LI indicates instability. Strong negative LI values support deep moist convection and severe updraft potential."),
    ("PWAT", "Precipitable Water. High PWAT supports heavy rainfall efficiency, stratiform persistence, and flood susceptibility when forcing is present."),
    ("SWEAT", "Severe Weather Threat Index. Combines instability, moisture, wind, and shear ingredients to indicate organized severe potential."),
    ("Bulk Shear", "Vector wind difference through a layer, commonly 0-6 km. Higher shear supports storm organization, propagation, and persistent cells."),
    ("LCL", "Lifted Condensation Level. Lower LCL indicates parcels reach saturation sooner, often supporting lower cloud bases and efficient convection."),
    ("LFC", "Level of Free Convection. The level where an air parcel becomes positively buoyant and can accelerate upward."),
    ("EL", "Equilibrium Level. The upper level where the parcel loses buoyancy; higher EL often implies taller storms and stronger echo tops."),
    ("Theta-E", "Equivalent potential temperature. High theta-e in low levels marks warm, moist air and instability corridors."),
    ("Thunderstorm Initiation", "The transition from conditional instability to active deep convection after lifting, convergence, heating, or boundary forcing overcomes inhibition."),
    ("Convective Instability", "A thermodynamic profile capable of supporting deep convection when adequate lift and moisture are present."),
    ("CSI", "Critical Success Index. Measures event forecast accuracy while penalizing misses and false alarms."),
    ("POD", "Probability of Detection. Fraction of observed events that were correctly forecast."),
    ("FAR", "False Alarm Ratio. Fraction of forecast events that did not occur; lower is better."),
    ("HSS", "Heidke Skill Score. Skill score relative to random chance, useful for categorical forecast verification."),
    ("BIAS", "Forecast frequency bias. Values greater than 1 indicate overforecasting; less than 1 indicates underforecasting."),
]


def md_table(headers: list[str], rows: list[list[str]]) -> str:
    def esc(value: str) -> str:
        return str(value).replace("\n", " ").replace("|", "\\|")
    out = ["| " + " | ".join(headers) + " |", "| " + " | ".join(["---"] * len(headers)) + " |"]
    for row in rows:
        out.append("| " + " | ".join(esc(x) for x in row) + " |")
    return "\n".join(out)


def code_block(lang: str, content: str) -> str:
    return f"```{lang}\n{content.strip()}\n```"


def module_profile(module: str) -> dict[str, str]:
    name = module.lower()
    if "doppler" in name or "radar" in name:
        return {
            "role": "Operational nowcasting forecaster",
            "workflow": "Select station, verify metadata HUD, inspect reflectivity-style markers, review threat summary, and continue monitoring or escalate.",
            "inputs": "Forecast cycle metadata, station focus, CAPE traceability, storm probability, lightning/rain/squall indicators.",
            "outputs": "Radar-like situational awareness, current hazard summary, station risk, and operational monitoring action.",
            "components": "Operational metadata HUD, station selector, radar map, threat summary, CAPE panel, risk panels.",
            "relevance": "Supports rapid visual assessment of thunderstorm, lightning, heavy-rain, and squall corridors.",
        }
    if "archive" in name or "historical" in name:
        return {
            "role": "IMD/CWC reviewer and research analyst",
            "workflow": "Select file/date/station, run historical pipeline, inspect latest events, registry, verification, and export evidence.",
            "inputs": "RSRW archive records, IMD observations, selected station/date, uploaded CSV/XLS/XLSX datasets.",
            "outputs": "Historical event evidence, latest thunderstorm, verification result, registry rows, and exportable reports.",
            "components": "Archive summary, file analysis center, selectors, latest event card, tables, review readiness panel.",
            "relevance": "Provides source-traceable investigation of historical thunderstorm records and forecast outcomes.",
        }
    if "forecast simulator" in name or "simulator" in name:
        return {
            "role": "Forecast reviewer and academic evaluator",
            "workflow": "Upload/paste/select sounding parameters, adjust station context, inspect validation stages, and export report.",
            "inputs": "CAPE, CIN, LI, PWAT, SWEAT, K-index, custom sounding text, or uploaded historical dataset.",
            "outputs": "Forecast class, probability, threshold trace, static-data warning, validation stages, and report exports.",
            "components": "Sliders, file input, validation stages, threshold comparison cards, forecast summary panels.",
            "relevance": "Shows how thermodynamic indices produce a reproducible thunderstorm forecast and warning recommendation.",
        }
    if "verification" in name:
        return {
            "role": "Verification scientist and threshold reviewer",
            "workflow": "Choose station/season/thresholds, recompute contingency metrics, inspect matrix and case match log.",
            "inputs": "Historical observations, forecast records, CAPE/LI/PWAT/SWEAT/K-index thresholds.",
            "outputs": "CSI, POD, FAR, BIAS, HSS, contingency matrix, case match log, and calibration advice.",
            "components": "Threshold controls, metric cards, matrix panels, case inspection table, operational advice panel.",
            "relevance": "Quantifies forecast skill and supports threshold calibration without changing forecast logic.",
        }
    if "dataset" in name or "file analysis" in name:
        return {
            "role": "Data engineer and reviewer",
            "workflow": "Upload dataset, preview records, detect columns, audit quality, generate registry, and export.",
            "inputs": "CSV, XLS, XLSX, radiosonde exports, historical sounding datasets.",
            "outputs": "Quality score, missing values, duplicates, station/date coverage, detected parameters, and registry.",
            "components": "Upload control, quality summary, station distribution, column audit, registry/export controls.",
            "relevance": "Confirms whether uploaded evidence is suitable for IMD/CWC historical analysis.",
        }
    if "reviewer" in name or "dashboard" in name:
        return {
            "role": "IMD/CWC reviewer",
            "workflow": "Open selected case, inspect forecast/verification evidence, enter reviewer verdict and remarks, export docket.",
            "inputs": "Selected historical case, verification result, meteorologist summary, reviewer fields.",
            "outputs": "Review docket, verdict, remarks, and CSV/XLSX/PDF exports.",
            "components": "Case audit panel, reviewer docket form, export controls, source-traceability summary.",
            "relevance": "Converts technical forecast evidence into an operational review record.",
        }
    if "coastal" in name or "nowcast" in name:
        return {
            "role": "Coastal severe weather monitoring desk",
            "workflow": "Inspect map, compare Visakhapatnam/Machilipatnam corridors, review inflow/convergence/probability metrics.",
            "inputs": "Station forecasts, map overlays, marine inflow, lightning/rainfall/squall probabilities.",
            "outputs": "Coastal corridor awareness, nowcast probability bars, and district monitoring guidance.",
            "components": "Map panel, station markers, probability bars, corridor metrics, coastal intelligence stream.",
            "relevance": "Supports north coastal Andhra and Bay of Bengal severe weather surveillance.",
        }
    return {
        "role": "Operational forecaster, reviewer, or evaluator",
        "workflow": "Open module, confirm metadata, inspect situation, reasoning, and recommended action.",
        "inputs": "Forecast state, station metadata, historical archive records, or selected workflow inputs.",
        "outputs": "Operational situation summary, meteorological reasoning, and recommended action.",
        "components": "Metadata HUD, metric cards, operational panels, tables, selectors, and export controls.",
        "relevance": "Provides reviewer-visible evidence within the StormSense AI operational workstation.",
    }


def screenshot_analysis_report() -> str:
    rows = []
    for item in SCREENSHOTS:
        profile = module_profile(item["module"])
        rows.append([
            item["id"],
            item["module"],
            item["caption"],
            profile["role"],
            profile["workflow"],
            profile["inputs"],
            profile["outputs"],
            profile["components"],
            profile["relevance"],
            "Keyboard and screen-reader review required for final WCAG certification; visual evidence shows labelled controls, readable status text, and operational hierarchy.",
            "Screens are dense operational workstations; evidence should be tested on laptop and desktop viewports for overflow and map/table rendering.",
            item["relative"],
        ])
    return "\n".join([
        "# Screenshot Analysis Report",
        "",
        f"Generated: {TODAY}",
        f"Screenshot evidence files analyzed: {len(SCREENSHOTS)}",
        "",
        "This report maps supplied screenshots to visible StormSense AI pages and reviewer workflows. The analysis is evidence-oriented and avoids fabricated page states.",
        "",
        md_table([
            "ID", "Page Name", "Purpose", "User Role", "Primary Workflow", "Inputs", "Outputs",
            "UI Components", "Operational Relevance", "Accessibility Notes", "Performance/Responsive Notes", "Local Evidence File"
        ], rows),
        "",
    ])


def requirements_specification() -> str:
    functional = [
        ["FR-01", "Authentication", "Allow users to log in and validate the active profile.", "`/auth/login`, `/auth/me`, `backend/main.py`", "Implemented"],
        ["FR-02", "Historical archive search", "Filter and inspect 2023-2025 thunderstorm records.", "`/cwc/historical-dates`, `ResearchHub.jsx`", "Implemented"],
        ["FR-03", "Forecast simulation", "Run station/custom sounding forecast reproduction.", "`/forecast`, `ResearchHub.jsx:FORECAST_LAB`", "Implemented"],
        ["FR-04", "File analysis", "Upload CSV/XLS/XLSX/radiosonde datasets and generate audit/registry output.", "`/cwc/analyze-historical-dataset`", "Implemented"],
        ["FR-05", "Verification", "Compute contingency metrics and threshold research.", "`/cwc/verification`, `/cwc/threshold-research`", "Implemented"],
        ["FR-06", "Dataset explorer", "Display archive coverage and CAPE dynamicity evidence.", "`ResearchHub.jsx:DATASET_EXPLORER`", "Implemented"],
        ["FR-07", "Reviewer dashboard", "Produce review docket, verdict, and export outputs.", "`ResearchHub.jsx:REVIEWER_DASHBOARD`", "Implemented"],
        ["FR-08", "Live radar/nowcast", "Display radar-like monitoring and coastal nowcast panels.", "`RadarConsole.jsx`, `Phase3OpsModule.jsx`", "Implemented"],
        ["FR-09", "Bulletin generation", "Generate operational bulletin products.", "`/cwc/operational-bulletins`", "Implemented"],
        ["FR-10", "Single URL deployment", "Serve React frontend and API from one Render service.", "`backend/main.py`, `render.yaml`", "Implemented"],
    ]
    nonfunctional = [
        ["NFR-01", "Reliability", "Backend should import safely under `uvicorn backend.main:app` and preserve API endpoints.", "Render deployment validation"],
        ["NFR-02", "Determinism", "Forecast, CAPE, probability, and verification outputs must remain deterministic and cycle-aware.", "Engine audit and endpoint regression"],
        ["NFR-03", "Security", "Secrets must be environment-based; no database credentials hardcoded in source.", ".env.example and deployment audit"],
        ["NFR-04", "Performance", "Operational pages should load without nested scroll failures and use API polling/websocket responsibly.", "Frontend build and browser review"],
        ["NFR-05", "Maintainability", "Frontend modules and backend engines must preserve existing ownership boundaries.", "Source reference matrix"],
        ["NFR-06", "Scalability", "Persistence layer should use Supabase PostgreSQL and SQLAlchemy connection pooling.", "Supabase migration reports"],
        ["NFR-07", "Availability", "Static frontend and API fallback should share one deployed service with SPA route refresh support.", "Single URL deployment report"],
        ["NFR-08", "Accessibility", "Operational labels, status text, and controls should remain readable and keyboard-auditable.", "Accessibility audit"],
    ]
    return "\n".join([
        "# Requirements Specification",
        "",
        f"Generated: {TODAY}",
        "",
        "## Functional Requirements",
        "",
        md_table(["ID", "Requirement", "Description", "Evidence", "Status"], functional),
        "",
        "## Non-Functional Requirements",
        "",
        md_table(["ID", "Category", "Requirement", "Validation"], nonfunctional),
        "",
        "## Traceability Note",
        "Requirements are derived from supplied screenshots, `frontend/src` modules, FastAPI route decorators, Supabase schema, and generated migration/deployment reports.",
        "",
    ])


def design_system_guide() -> str:
    return "\n".join([
        "# StormSense AI Design System Guide",
        "",
        f"Generated: {TODAY}",
        "",
        "## Design Intent",
        "StormSense AI uses an operational workstation style: dark panels, high-contrast status text, restrained accent colors, dense but readable metadata, and command-deck navigation. This guide documents the existing interface; it does not propose a redesign.",
        "",
        "## Typography",
        md_table(["Use", "Font"], [
            ["Headings, navigation, labels, buttons, cards, table headers", "Satoshi"],
            ["Dates, times, station codes, metadata, parameters, verification codes, forecast IDs", "JetBrains Mono"],
            ["Fallback stack", "Inter, Roboto, system sans-serif where configured by the frontend"],
        ]),
        "",
        "## Interface Tokens",
        md_table(["Element", "Observed Pattern"], [
            ["Page shell", "Dark operational canvas with fixed sidebar and content workspace."],
            ["Cards/panels", "Low-radius dark panels with subtle borders and status accent colors."],
            ["Buttons", "Compact action buttons, primary blue for action, accent red/amber/green/purple for status."],
            ["Inputs", "Dark select/input fields with visible borders and mono values where technical."],
            ["Tables", "Dense operational rows with sticky/visible headers where useful and overflow protection."],
            ["Accordions", "Advanced/scientific details collapsed by default where complexity is high."],
            ["Metadata HUD", "Observation date/time, forecast generated, station, cycle, and data source in a compact mono layout."],
        ]),
        "",
        "## Status Color Semantics",
        md_table(["Color Family", "Operational Meaning"], [
            ["Blue", "Primary action, selected station, forecast metadata, stable operation."],
            ["Green", "Online, accepted, pass, source-traceable, favorable completed state."],
            ["Amber/Yellow", "Monitoring, caution, heavy rain, cycle timing, intermediate risk."],
            ["Red/Pink", "Critical alert, miss, static data warning, severe thunderstorm risk."],
            ["Purple", "Review/evidence accent, severe escalation, advanced scientific detail."],
        ]),
        "",
        "## Screenshot Evidence",
        f"{len(SCREENSHOTS)} supplied screenshots are catalogued in `SCREENSHOT_ANALYSIS_REPORT.md` and copied under `docs/screenshots/`.",
        "",
    ])


def accessibility_audit() -> str:
    rows = [
        ["Keyboard navigation", "Primary controls are standard buttons/selects/inputs in React modules.", "Manual keyboard traversal still required before public release."],
        ["ARIA semantics", "Operational status text is visible; icon-only controls should have accessible labels where implemented.", "Audit with browser accessibility tree recommended."],
        ["Color contrast", "Screenshots show high-contrast text on dark panels and color-coded status labels.", "Run automated WCAG contrast check on production URL."],
        ["Responsive layouts", "Screenshots include laptop/desktop evidence; archive/table overflow issues were previously targeted.", "Validate 1366x768, 1600x900, 1920x1080."],
        ["Focus states", "Buttons and inputs are visible; focus state needs browser validation.", "Use keyboard-only review mode walkthrough."],
        ["Screen readers", "Complex maps/charts require textual companion summaries.", "Ensure maps/charts retain adjacent operational summary text."],
        ["Review Mode", "Reduces visual complexity and hides diagnostic clutter.", "Validate reviewer can understand a case within 20-30 seconds."],
    ]
    return "\n".join([
        "# Accessibility Audit",
        "",
        f"Generated: {TODAY}",
        "",
        "This audit records evidence visible in source/screenshots and identifies the final checks needed for WCAG 2.1 AA certification. It does not certify WCAG compliance by itself.",
        "",
        md_table(["Area", "Evidence", "Remaining Validation"], rows),
        "",
    ])


def security_guide() -> str:
    rows = [
        ["Authentication", "`/auth/login`, `/auth/signup`, `/auth/me` in FastAPI.", "Validate password hashing policy and token expiry in deployment."],
        ["Authorization", "Bearer-token flow used by authenticated profile checks.", "Add route-level policy audit for production roles."],
        ["Secrets", "Database/Supabase credentials are environment variables; `.env.example` documents required keys.", "Never commit `.env` values."],
        ["Database", "Supabase PostgreSQL schema and SQLAlchemy connection layer.", "Review RLS/Supabase policies before external launch."],
        ["File uploads", "File analysis endpoint supports CSV/XLS/XLSX/radiosonde datasets.", "Limit file size, validate MIME/type, sanitize names, and scan inputs as deployment policy."],
        ["CORS", "CORS allow-list can be set through environment policy.", "Restrict production origins."],
        ["Audit logs", "Migration/auth audit tables exist in schema.", "Verify log retention and PII policy."],
    ]
    return "\n".join([
        "# Security Guide",
        "",
        f"Generated: {TODAY}",
        "",
        md_table(["Security Area", "Current Evidence", "Production Control"], rows),
        "",
        "## Deployment Rule",
        "Use environment variables only for database URLs, Supabase keys, JWT secrets, and deployment origins. Do not hardcode credentials in frontend or backend source files.",
        "",
    ])


def testing_report() -> str:
    rows = [
        ["Build", "`npm run build`", "Generate `frontend/dist` for single-URL Render deployment."],
        ["Backend import", "`python -c \"import backend.main\"`", "No `ModuleNotFoundError` under package startup."],
        ["API smoke", "`/system-status`, `/history`, `/forecast`, `/cwc/historical-dates`", "HTTP 200 JSON response and stable schema."],
        ["Websocket", "`/stream/atmospheric`", "Receive operational telemetry frame or graceful fallback."],
        ["Upload", "`/cwc/analyze-historical-dataset`", "Quality summary, station/parameter detection, registry, and export metadata."],
        ["Archive", "Historical Thunderstorm Archive UI", "2023, 2024, 2025 records participate in summaries and registry."],
        ["Verification", "Threshold verification lab", "CSI, POD, FAR, HSS, BIAS computed from historical observations."],
        ["Review Mode", "Header Review Mode toggle", "Reviewer-critical fields visible; diagnostics hidden."],
        ["Deployment", "Render single URL", "Root and browser routes serve React shell; APIs remain active."],
    ]
    return "\n".join([
        "# Testing Report",
        "",
        f"Generated: {TODAY}",
        "",
        md_table(["Test Area", "Method", "Acceptance Condition"], rows),
        "",
        testing_details_for_main(),
        "",
    ])


def screenshot_evidence_section() -> str:
    if not SCREENSHOTS:
        return "\n".join([
            "## Screenshot Evidence Catalogue",
            "",
            "No local screenshot files were found under `docs/screenshots/` at generation time.",
            "",
        ])
    rows = [[item["id"], item["module"], item["caption"], item["relative"]] for item in SCREENSHOTS]
    selected = SCREENSHOTS[:SELECTED_SCREENSHOT_LIMIT]
    text = [
        "## Screenshot Evidence Catalogue",
        "",
        f"Supplied screenshot evidence files copied into `docs/screenshots/`: {len(SCREENSHOTS)}.",
        "",
        md_table(["ID", "UI Area", "Evidence Summary", "Local File"], rows),
        "",
        "### Embedded Evidence Figures",
        "",
    ]
    for item in selected:
        text += [
            f"#### {item['id']} - {item['module']}",
            item["caption"],
            "",
            f"![{item['id']} {item['module']}]({item['relative']})",
            "",
        ]
    if len(SCREENSHOTS) > len(selected):
        text.append(f"Additional screenshots are catalogued above and retained in `docs/screenshots/`; the main dossier embeds the first {len(selected)} representative figures to keep the PDF/DOCX readable.")
        text.append("")
    return "\n".join(text)


def mvp_documentation_version() -> str:
    selected = SCREENSHOTS[:8]
    text = [
        "# StormSense AI MVP Documentation Version",
        "",
        f"Generated: {TODAY}",
        "",
        "## Purpose",
        "This MVP version gives reviewers a compact first-read package before opening the full dossier. It focuses on what the platform is, how it is used, what evidence exists, and which workflows should be demonstrated.",
        "",
        "## Executive Summary",
        f"StormSense AI is a FastAPI + Vite React operational thunderstorm decision-support workstation with {len(ROUTES)} backend routes, {len(SCHEMA)} Supabase tables, {COUNTS['historical_archive_records']} historical archive records, {COUNTS['observational_records']} observational records, and {len(SCREENSHOTS)} supplied screenshot evidence files.",
        "",
        "## Core Operational Workflows",
        md_table(["Workflow", "Reviewer Action", "Evidence"], [
            ["Live Doppler/Radar Monitoring", "Open Live Doppler Radar, confirm station and cycle metadata, inspect risk summary.", "`RadarConsole.jsx`, screenshot evidence IMG-24"],
            ["Historical Archive", "Open Historical Thunderstorm Archive, inspect latest event, archive counts, registry, and selected case.", "`ResearchHub.jsx`, `/cwc/historical-dates`, IMG-33/IMG-45"],
            ["Forecast Simulator", "Run or review custom sounding simulation and validation stages.", "`/forecast`, `ResearchHub.jsx:FORECAST_LAB`, IMG-34/IMG-43"],
            ["File Analysis Center", "Upload historical dataset and review quality/registry output.", "`/cwc/analyze-historical-dataset`, IMG-23"],
            ["Verification", "Recompute threshold scores and inspect contingency matrix.", "`/cwc/verification`, IMG-21/IMG-22"],
            ["Reviewer Dashboard", "Record case verdict and export review docket.", "`ResearchHub.jsx:REVIEWER_DASHBOARD`, IMG-36"],
        ]),
        "",
        "## System Architecture Snapshot",
        code_block("mermaid", """
flowchart LR
  UI["React Operational Workstation"] --> API["FastAPI Backend"]
  API --> Engines["Forecast, Thermodynamic, Verification, Archive Engines"]
  API --> DB["Supabase PostgreSQL"]
  API --> Uploads["Dataset Upload and Export Workflows"]
  Engines --> Review["Reviewer Evidence and IMD/CWC Outputs"]
        """),
        "",
        "## Evidence Figures",
        "",
    ]
    for item in selected:
        text += [
            f"### {item['id']} - {item['module']}",
            item["caption"],
            f"Local file: `{item['relative']}`",
            "",
        ]
    text += [
        "## Readiness Notes",
        "- The full dossier contains complete API, database, design system, accessibility, security, testing, user, developer, and deployment documentation.",
        "- The authoritative refreshed PDF is `StormSense_AI_IMD_Documentation_v2_1.pdf` because the older PDF filename was locked by another process during regeneration.",
        "",
    ]
    return "\n".join(text)


def full_documentation_expansion_plan() -> str:
    rows = [
        ["1", "Evidence Capture", "Screenshot catalogue, source inventory, route/schema extraction", "Complete", "Developer + reviewer"],
        ["2", "Core Dossier", "Full IMD/CWC operational documentation package", "Complete", "Developer"],
        ["3", "MVP Brief", "Compact reviewer-first documentation version", "Complete", "Developer"],
        ["4", "Operations Appendix", "Add live review notes from internal testing", "Pending review session", "IMD/CWC reviewer"],
        ["5", "Accessibility Certification", "Browser accessibility tree, keyboard traversal, WCAG contrast report", "Needs manual/browser audit", "UX/accessibility reviewer"],
        ["6", "Security Certification", "RLS policy review, secret rotation, upload policy, auth audit", "Needs production policy confirmation", "Security owner"],
        ["7", "Deployment Evidence", "Render production logs, endpoint screenshots, Supabase row-count snapshots", "After deployment window", "Platform engineer"],
    ]
    return "\n".join([
        "# Full Documentation Expansion Plan",
        "",
        f"Generated: {TODAY}",
        "",
        "This plan keeps the submitted package production-oriented while identifying evidence that should be appended after formal internal testing.",
        "",
        md_table(["Phase", "Focus", "Deliverable", "Status", "Owner"], rows),
        "",
        "## Final Package Baseline",
        "The complete package currently includes Markdown, DOCX, and PDF dossier outputs plus standalone screenshot, requirements, design, accessibility, security, testing, API, user, developer, architecture, and deployment documents.",
        "",
        "## Next Review Additions",
        "- Add internal test run IDs only if the review body explicitly requests them. This generated package intentionally uses document-control front matter only.",
        "- Add production URL screenshots after Render deployment is live and stable.",
        "- Add Supabase dashboard evidence after final row-count validation.",
        "- Add accessibility certification artifacts after browser-based WCAG review.",
        "",
    ])


def architecture_diagrams() -> str:
    diagrams = [
        ("High-Level Architecture", """
flowchart LR
  Operator["IMD/CWC Reviewer or Forecaster"] --> UI["Vite React Workstation"]
  UI --> API["FastAPI Operational Backend"]
  API --> Sounding["Wyoming / Cached Sounding Ingestion"]
  API --> Engines["Thermodynamic, Forecast, Verification, Research Engines"]
  API --> DB["Supabase PostgreSQL"]
  API --> Files["CSV/XLS/XLSX Upload and Export Workflows"]
  Engines --> Outputs["Nowcast, Verification, Registry, Bulletins, District Impacts"]
        """),
        ("Component Architecture", """
flowchart TB
  App["frontend/src/App.jsx"] --> Sidebar["Sidebar Navigation"]
  App --> Radar["RadarConsole"]
  App --> Predictor["PredictorEngine"]
  App --> Analytics["AnalyticsDeck"]
  App --> Research["ResearchHub"]
  App --> Phase3["Phase3OpsModule"]
  App --> Health["HealthConsole"]
  Research --> Archive["Historical Thunderstorm Archive"]
  Research --> Simulator["Forecast Simulator"]
  Research --> Verify["Forecast Verification"]
  Phase3 --> Andhra["Coastal Thunderstorm Monitoring Center"]
  Phase3 --> Bulletin["Auto IMD Bulletin Generator"]
        """),
        ("Deployment Architecture", """
flowchart LR
  Render["Single Render Web Service"] --> Build["Build: pip install + npm install + npm run build"]
  Build --> Dist["frontend/dist"]
  Render --> Uvicorn["uvicorn backend.main:app"]
  Uvicorn --> Static["FastAPI serves /assets and SPA fallback"]
  Uvicorn --> API["FastAPI API routes"]
  API --> Supabase["Supabase PostgreSQL"]
        """),
        ("Authentication Flow", """
sequenceDiagram
  participant User
  participant React
  participant FastAPI
  participant DB as Supabase users table
  User->>React: Enter login credentials
  React->>FastAPI: POST /auth/login
  FastAPI->>DB: Lookup user by email
  DB-->>FastAPI: User record
  FastAPI-->>React: Bearer token and profile
  React->>FastAPI: Authenticated requests with Authorization header
  FastAPI-->>React: Protected operational outputs
        """),
        ("Forecast Pipeline", """
flowchart TD
  Cycle["00Z / 12Z Cycle Lock"] --> Sounding["Sounding Fetch and Cache"]
  Sounding --> Thermo["Thermodynamic Solver"]
  Thermo --> Indices["CAPE, CIN, LI, PWAT, SWEAT, K-index"]
  Indices --> Prob["Probabilistic Forecast Engine"]
  Indices --> Lifecycle["Convective Lifecycle Engine"]
  Prob --> Decision["Operational Decision Support"]
  Lifecycle --> Decision
  Decision --> Bulletin["Bulletin / Watch / Advisory Outputs"]
  Decision --> Verification["Verification and Trust Scoring"]
        """),
    ]
    text = ["# Architecture Diagrams", "", f"Generated from repository inspection on {TODAY}.", ""]
    for title, diagram in diagrams:
        text += [f"## {title}", "", code_block("mermaid", diagram), ""]
    return "\n".join(text)


def api_reference() -> str:
    text = ["# StormSense AI API Reference", "", f"Generated from `backend/main.py` on {TODAY}.", ""]
    text.append(f"Total API/websocket route decorators discovered: {len(ROUTES)}.")
    text.append("")
    rows = [[r["method"], r["path"], r["handler"], route_purpose(r), f"backend/main.py:{r['line']}"] for r in ROUTES]
    text.append(md_table(["Method", "Route", "Handler", "Purpose", "Source"], rows))
    text.append("")
    text.append("## Endpoint Notes")
    for r in ROUTES:
        if r["path"] == "/{full_path:path}":
            continue
        auth = "Bearer token required" if r["path"].startswith("/auth/me") or r["path"] in ["/cwc/override", "/cwc/clear-override", "/cwc/cycle"] and r["method"] == "POST" else "Public operational API or UI-facing endpoint"
        text += [
            f"### {r['method']} {r['path']}",
            f"Purpose: {route_purpose(r)}",
            f"Handler: `{r['handler']}` at `backend/main.py:{r['line']}`.",
            f"Authentication: {auth}.",
            "Input: query parameters, JSON body, uploaded file, or websocket frame according to the route definition.",
            "Output: JSON operational payload unless the route is a websocket, CSV export, raw sounding, or SPA static response.",
            "Error handling: FastAPI status responses, route-specific validation, and fallback cache behavior where implemented.",
            "",
        ]
    return "\n".join(text)


def user_manual() -> str:
    sections = [
        ("Login", ["Open the single Render URL.", "Enter reviewer/operator credentials.", "Confirm the operational HUD shows backend and websocket status.", "Use guest/local state only for non-production demonstration when authorized."]),
        ("Run Forecast", ["Open AI Prediction Engine or Thunderstorm Forecast Simulator.", "Confirm station, cycle, and data source metadata.", "Review CAPE, LI, PWAT, SWEAT, K-index, forecast result, probability, and recommended action.", "Use Forecast Generated time and cycle lock to verify operational timing."]),
        ("Analyze Historical Case", ["Open Research & Insights.", "Select Historical Thunderstorm Archive.", "Filter by date, station, season, and event type.", "Open historical analysis to compare forecast reproduction with observed event."]),
        ("Upload Dataset", ["Open Historical Thunderstorm Archive or Thunderstorm Forecast Simulator.", "Use Analyze Historical Dataset.", "Upload CSV, XLS, XLSX, radiosonde, or historical sounding file.", "Review dataset quality summary, detected columns, registry, verification, and export options."]),
        ("Review Verification", ["Open Forecast Verification.", "Select station, season, thresholds, or historical case.", "Review CSI, HSS, POD, FAR, and BIAS.", "Use the results to decide whether thresholds require calibration."]),
        ("Generate Bulletin", ["Open Auto IMD Bulletin Generator.", "Review current station hazards and district impact logic.", "Generate thunderstorm, lightning, heavy rainfall, coastal squall, or district summary outputs.", "Export or copy the bulletin for reviewer discussion."]),
        ("Export Reports", ["Use CSV, JSON, XLSX, PDF, or analysis export controls where available.", "Confirm source dataset and station metadata.", "Attach reviewer docket metadata when required."]),
        ("Use Review Mode", ["Toggle IMD Review Mode from the header.", "Confirm charts and diagnostics are hidden.", "Review date, time, station, observed event, forecast result, verification result, meteorologist summary, and recommended action."]),
    ]
    text = ["# StormSense AI User Manual", "", "Audience: IMD/CWC reviewer, operational forecaster, academic evaluator, and demonstration presenter.", ""]
    for title, steps in sections:
        text += [f"## {title}", ""]
        for idx, step in enumerate(steps, 1):
            text.append(f"{idx}. {step}")
        text.append("")
    return "\n".join(text)


def api_details_for_main() -> str:
    text = ["## Complete API Endpoint Detail", ""]
    for r in ROUTES:
        if r["path"] == "/{full_path:path}":
            continue
        path = r["path"]
        method = r["method"]
        purpose = route_purpose(r)
        auth = "Bearer token required" if path.startswith("/auth/me") or (path in ["/cwc/override", "/cwc/clear-override", "/cwc/cycle"] and method == "POST") else "No bearer token required for standard dashboard/review retrieval unless deployment policy adds external access control."
        if method == "POST" and "upload" in path or "analyze-historical-dataset" in path:
            input_text = "Multipart file upload. Supported operational dataset formats include CSV, XLS, XLSX, radiosonde text, and historical sounding datasets where implemented."
        elif method == "POST":
            input_text = "JSON request body matching the route's Pydantic model or route-specific payload fields."
        elif method == "WEBSOCKET":
            input_text = "Websocket connection request from the frontend telemetry client."
        else:
            input_text = "Optional query parameters where declared by the route definition; otherwise no request body."
        output_text = "JSON response containing operational data, metadata, diagnostics, or verification results."
        if "export" in path:
            output_text = "CSV, JSON, or analysis export response depending on route and query parameters."
        if "sounding-raw" in path:
            output_text = "Raw sounding text plus sounding metadata for source traceability."
        if method == "WEBSOCKET":
            output_text = "Streaming JSON frames containing forecasts, trends, escalations, and cycle metadata."
        text += [
            f"### {method} {path}",
            f"Source reference: `backend/main.py:{r['line']}` handler `{r['handler']}`.",
            f"Purpose: {purpose}",
            f"Input: {input_text}",
            f"Output: {output_text}",
            f"Authentication: {auth}",
            "Operational review note: validate this endpoint during IMD/CWC review if its paired frontend module is demonstrated. Confirm HTTP status, payload shape, source metadata, and user-visible module behavior.",
            "",
        ]
    return "\n".join(text)


def database_details_for_main() -> str:
    text = ["## Table-by-Table Supabase Notes", ""]
    descriptions = {
        "roles": "Role catalog for operational user classification.",
        "users": "Persisted application users migrated from the MySQL Workbench export and used by FastAPI authentication.",
        "profiles": "Optional profile extension table linked to users and future Supabase auth identifiers.",
        "audit_login": "Login audit trail recording successful and failed authentication attempts.",
        "thunderstorm_forecasts": "Primary forecast persistence table containing station, indices, forecast class, probability, and created timestamp.",
        "historical_observations": "Normalized historical observation records used by archive and verification workflows.",
        "historical_records": "Historical record mirror/compatibility table for archive workflows.",
        "thunderstorm_registry": "File/archive derived event registry with thunderstorm, NWX, severe storm, forecast, verification, and season flags.",
        "verification_results": "Forecast verification outputs, including CSI, HSS, POD, FAR, and BIAS.",
        "uploaded_datasets": "File Analysis Center persistence for upload metadata, quality score, and analysis JSON.",
        "reviewer_dashboard_records": "Reviewer docket, verdict, comments, station, date, and export metadata.",
        "audit_logs": "General operational audit log for migration and backend activity.",
    }
    for table, cols in SCHEMA.items():
        text += [
            f"### {table}",
            descriptions.get(table, "Supabase persistence table discovered in `supabase_schema.sql`."),
            "",
            md_table(["Column Definition"], [[col] for col in cols]),
            "",
            "Review validation: confirm row count after migration, primary key behavior, nullable fields, and related frontend workflow visibility.",
            "",
        ]
    return "\n".join(text)


def testing_details_for_main() -> str:
    scenarios = [
        ("Authentication Testing", "POST `/auth/login` with a known reviewer account, verify token issuance, store token in frontend local storage, call `/auth/me`, then logout and confirm protected actions are unavailable."),
        ("Historical Archive Testing", "Open Research & Insights -> Historical Thunderstorm Archive. Confirm 2023, 2024, and 2025 records are visible in summaries, latest thunderstorm calculation, registry filters, and historical analysis."),
        ("File Upload Testing", "Upload CSV/XLS/XLSX files through Analyze Historical Dataset. Confirm station detection, parameter detection, missing values, duplicate rows, quality score, registry generation, and latestFileAnalyzed update."),
        ("Forecast Simulator Testing", "Run the simulator for Visakhapatnam and Machilipatnam. Confirm deterministic outputs, threshold trace, probability explanation, forecast result, and verification result."),
        ("Verification Testing", "Run threshold research by station/season and verify CSI, HSS, POD, FAR, and BIAS change deterministically with threshold input."),
        ("CAPE Traceability Testing", "Call `/cwc/cape-traceability`, check NOW through T-6 timeline, source status, cache age, delta CAPE, and STATIC_DATA_WARNING behavior."),
        ("SPA Deployment Testing", "Open Render URL at `/`, `/dashboard`, `/research`, `/archive`, and `/forecast`. Refresh each route and confirm React shell loads."),
        ("API Preservation Testing", "Call `/system-status`, `/history`, `/forecast`, `/cwc/historical-dates`, and `/cwc/historical-analysis`. Confirm JSON payloads remain active after static frontend serving."),
        ("Websocket Testing", "Connect to `/stream/atmospheric` and verify forecast/trend/escalation/cycle frames arrive or fall back cleanly."),
        ("Reviewer Acceptance Testing", "Enable IMD Review Mode and confirm the reviewer sees date, time, station, observed event, forecast result, verification result, meteorologist summary, and recommended action without diagnostics clutter."),
    ]
    text = ["## Detailed Test Scenarios", ""]
    for title, detail in scenarios:
        text += [f"### {title}", detail, "Expected outcome: pass/fail result must be recorded with endpoint, UI location, dataset/source, and reviewer-visible evidence.", ""]
    return "\n".join(text)


def developer_manual() -> str:
    rows = [
        ["Frontend entry", "frontend/src/App.jsx", "Authentication, global telemetry state, websocket connection, and workstation module routing."],
        ["Sidebar", "frontend/src/components/layout/Sidebar.jsx", "Primary navigation module list."],
        ["Research center", "frontend/src/components/modules/ResearchHub.jsx", "Start Here, About, archive, simulator, verification, dataset explorer, reviewer dashboard, index guide, terminology."],
        ["Phase 3 ops modules", "frontend/src/components/modules/Phase3OpsModule.jsx", "Atmospheric intelligence, watchdesk, lab, bulletin, analog, Andhra monitoring, verification."],
        ["Backend app", "backend/main.py", "FastAPI routes, authentication, websocket, forecasting, archive endpoints, file upload endpoints, static SPA serving."],
        ["Research engines", "backend/analysis_engines.py", "Historical observations, verification, probability, climatology, analog, district impact, bulletins."],
        ["Thermodynamics", "backend/thermo.py", "Sounding parsing and thermodynamic solver."],
        ["Sounding ingestion", "backend/fetch_sounding.py", "Wyoming fetch, cache metadata, freshness scoring, cycle-aware cache."],
        ["Database layer", "backend/database.py and backend/connection_pool.py", "SQLAlchemy PostgreSQL connection and Supabase schema initialization helpers."],
    ]
    text = ["# StormSense AI Developer Manual", "", "## Project Structure", "", md_table(["Area", "Path", "Responsibility"], rows), ""]
    text += [
        "## Local Development",
        "1. Install backend requirements from `backend/requirements.txt`.",
        "2. Install frontend dependencies from `frontend/package.json`.",
        "3. Set `.env` values for Supabase/PostgreSQL and authentication secrets.",
        "4. Run FastAPI with `uvicorn backend.main:app` from the repository root.",
        "5. Run `npm run dev` in `frontend/` only for local frontend development; production uses FastAPI static serving.",
        "",
        "## Code Organization Rules",
        "- Keep scientific functions in `analysis_engines.py`, `thermo.py`, and `fetch_sounding.py` deterministic.",
        "- Keep UI refinements inside existing components unless a component already owns the workflow.",
        "- Do not hardcode deployment URLs in frontend code; production resolves same-origin through `environment.js`.",
        "- Preserve package-safe backend imports because Render starts `uvicorn backend.main:app`.",
        "",
        "## Database Migration",
        "- Supabase schema is maintained in `supabase_schema.sql`.",
        "- Migration execution is handled by `migration_runner.py`.",
        "- Runtime database access uses SQLAlchemy through `backend/database.py` and `backend/connection_pool.py`.",
        "",
        "## Contribution Guidelines",
        "- Run frontend build before deployment handoff.",
        "- Run backend import and compile validation after backend changes.",
        "- Preserve CAPE, probability, verification, and archive determinism.",
        "- Add documentation evidence for any new operational module or endpoint.",
    ]
    return "\n".join(text)


def deployment_guide() -> str:
    return "\n".join([
        "# StormSense AI Deployment Guide",
        "",
        "## Target Deployment",
        "StormSense AI is configured for a single Render Web Service serving both FastAPI APIs and the Vite React frontend from one URL.",
        "",
        "## Render Build Process",
        code_block("text", """
python -m pip install --upgrade pip
pip install -r backend/requirements.txt
cd frontend && npm install && npm run build
        """),
        "",
        "## Render Start Command",
        code_block("text", "uvicorn backend.main:app --host 0.0.0.0 --port $PORT"),
        "",
        "## Environment Variables",
        md_table(["Variable", "Purpose"], [
            ["DATABASE_URL", "Supabase PostgreSQL connection string with SSL mode."],
            ["SUPABASE_URL", "Supabase project API URL."],
            ["SUPABASE_ANON_KEY", "Client-safe Supabase key where needed."],
            ["SUPABASE_SERVICE_ROLE_KEY", "Server-side Supabase service key; never expose in frontend."],
            ["JWT_SECRET", "Authentication signing secret."],
            ["CORS_ALLOW_ORIGINS", "Optional comma-separated CORS allow-list."],
        ]),
        "",
        "## Static Frontend Serving",
        "FastAPI serves `frontend/dist/assets` at `/assets` and returns `frontend/dist/index.html` for browser navigation routes.",
        "",
        "## API Preservation",
        "Backend prefixes `/auth`, `/cwc`, `/history`, `/system-status`, `/stream`, `/trend-analysis`, and `/storm-escalation` are protected from SPA fallback.",
        "",
        "## Troubleshooting",
        "- If the Render URL returns backend JSON at `/`, confirm the frontend build exists before startup.",
        "- If `/forecast` opens JSON in a browser, confirm the browser sends `Accept: text/html`; API calls use `Accept: application/json`.",
        "- If uploads fail, verify `python-multipart`, `openpyxl`, and `xlrd` are installed from requirements.",
        "- If database writes fail, verify `DATABASE_URL` and Supabase network access.",
    ])


def page_docs() -> str:
    pages = [
        ("Dashboard / Operational Shell", "frontend/src/App.jsx", "Authenticated landing workspace that coordinates HUD status, station forecasts, websocket telemetry, and module routing.", "Forecaster or reviewer.", "Login token, forecast API payloads, websocket stream, cycle info.", "Operational module views, HUD status, active cycle, and navigation context."),
        ("Live Doppler Radar", "frontend/src/components/modules/RadarConsole.jsx", "Radar-like operational visualization driven by forecast data, lifecycle state, lightning, rainfall, and station risk.", "Nowcasting forecaster.", "Forecast rows, storm probability, station metadata.", "Operational radar console, severe markers, and monitoring interpretation."),
        ("AI Prediction Engine", "frontend/src/components/modules/PredictorEngine.jsx", "Station-level deterministic forecast workstation with convective index interpretation.", "Forecast desk reviewer.", "CAPE, LI, SWEAT, PWAT, K-index, cycle data.", "Forecast classification, probability, trend, and action guidance."),
        ("Convective Analytics", "frontend/src/components/modules/AnalyticsDeck.jsx", "Verification, analytics, climatology, and decision-support panels.", "Research analyst and verification reviewer.", "Historical observations, thresholds, forecast rows.", "Skill metrics, reliability, threshold insights, and forecast trust context."),
        ("Research & Insights", "frontend/src/components/modules/ResearchHub.jsx", "Multi-tab research center containing archive, simulator, verification, database, review, glossary, and onboarding pages.", "Reviewer, researcher, and presenter.", "Historical archive, upload file, selected station/date, thresholds.", "Case analysis, registry, verification, dataset audit, and reviewer docket."),
        ("Historical Thunderstorm Archive", "frontend/src/components/modules/ResearchHub.jsx:HISTORICAL_WORKBENCH", "Flagship meteorological investigation console for historical thunderstorm records.", "IMD/CWC reviewer.", "2023-2025 archive records, file upload, station/date filters.", "Latest records, thunderstorm registry, archive summaries, exports, and historical analysis."),
        ("Thunderstorm Forecast Simulator", "frontend/src/components/modules/ResearchHub.jsx:FORECAST_LAB", "Historical and custom sounding forecast reproduction lab.", "Forecaster and academic evaluator.", "Station, date, custom sounding values, uploaded dataset.", "Simulated forecast, threshold trace, probability, verification, and exportable report."),
        ("Forecast Verification", "frontend/src/components/modules/ResearchHub.jsx:RESEARCH_VERIFY", "Threshold testing and contingency-matrix verification workspace.", "Verification scientist.", "Thresholds, station, season, historical records.", "CSI, POD, FAR, HSS, BIAS, and recommended calibration direction."),
        ("Dataset Explorer", "frontend/src/components/modules/ResearchHub.jsx:DATASET_EXPLORER", "Historical weather database browser.", "Research analyst.", "Archive index and selected station/date filters.", "Weather archive log and searchable historical observations."),
        ("IMD Review Dashboard", "frontend/src/components/modules/ResearchHub.jsx:REVIEWER_DASHBOARD", "Reviewer docket and operational review console.", "Reviewer or evaluator.", "Selected event, reviewer name, docket ID, comments.", "Review summary, verdict, action, export metadata."),
        ("About StormSense AI", "frontend/src/components/modules/ResearchHub.jsx:ABOUT", "Operational reference page for purpose, workflow, data sources, architecture, and outputs.", "Reviewer and evaluator.", "Static project metadata and workflow descriptions.", "Architecture and operational purpose narrative."),
        ("Start Here", "frontend/src/components/modules/ResearchHub.jsx:START_HERE", "Reviewer onboarding center.", "First-time reviewer.", "None beyond app state.", "Data requirements, outputs, common review flow, and quick navigation."),
        ("File Analysis Center", "frontend/src/components/modules/ResearchHub.jsx and Phase3OpsModule.jsx", "Upload and audit historical datasets.", "Research analyst and reviewer.", "CSV, XLS, XLSX, radiosonde, and sounding datasets.", "Quality summary, column detection, registry, verification, and report exports."),
        ("Operational Monitoring", "frontend/src/components/modules/Phase3OpsModule.jsx", "Severe watchdesk, Andhra monitoring, radar-sounding fusion, and operational intelligence modules.", "Nowcasting desk.", "Forecast data, lifecycle state, station focus, district selection.", "Monitoring recommendation, district impacts, and severe-weather guidance."),
        ("Bulletin Generator", "frontend/src/components/modules/Phase3OpsModule.jsx:BULLETIN", "Operational bulletin generation workspace.", "Forecaster preparing advisory text.", "Current forecast rows and cycle metadata.", "Thunderstorm nowcasts, lightning advisories, heavy rainfall bulletins, and district summaries."),
        ("Health & Migrations", "frontend/src/components/modules/HealthConsole.jsx", "Deployment/runtime health and migration visibility console.", "Developer/operator.", "API base URL, websocket URL, system status.", "Backend health, connectivity, and deployment readiness status."),
    ]
    rows = []
    for name, source, purpose, users, inputs, outputs in pages:
        rows.append([name, source, purpose, users, inputs, outputs])
    text = [
        "## Page-by-Page Documentation",
        "",
        md_table(["Page", "Source Reference", "Purpose", "Primary Users", "Inputs", "Outputs"], rows),
        "",
        f"Screenshot evidence note: {len(SCREENSHOTS)} supplied UI screenshots are copied under `docs/screenshots/` and catalogued in `SCREENSHOT_ANALYSIS_REPORT.md`.",
        "",
    ]
    for name, source, purpose, users, inputs, outputs in pages:
        text += [
            f"### {name}",
            f"Source reference: `{source}`.",
            f"Purpose: {purpose}",
            f"Who uses it: {users}",
            f"Inputs: {inputs}",
            f"Outputs: {outputs}",
            "Workflow: open the page from the sidebar or Research & Insights tab, confirm the operational metadata HUD, select station/date/cycle or file where applicable, review the current situation, meteorological interpretation, and recommended action, then export or record the reviewer verdict when the workflow supports it.",
            "Operational relevance: this page contributes to the decision chain by answering what is happening, why it is happening, what may happen next, and what operational action should be taken.",
            "Example usage: during IMD/CWC review, demonstrate the page with Visakhapatnam and Machilipatnam focus where station selection is available, then connect the visible output to the relevant backend endpoint or archive record.",
            "Screenshot reference: see the screenshot evidence catalogue and `SCREENSHOT_ANALYSIS_REPORT.md` for matching supplied UI evidence.",
            "",
        ]
    return "\n".join(text)


def main_document() -> str:
    route_rows = [[r["method"], r["path"], r["handler"], f"backend/main.py:{r['line']}"] for r in ROUTES]
    sidebar_rows = [[item["id"], item["label"]] for item in SIDEBAR]
    schema_rows = [[name, ", ".join(col.split()[0] for col in cols[:8]) + (" ..." if len(cols) > 8 else "")] for name, cols in SCHEMA.items()]
    research_rows = [[code, label, desc] for code, label, desc in RESEARCH_TABS]
    glossary_rows = [[term, desc] for term, desc in GLOSSARY]

    text = [
        "# StormSense AI IMD/CWC Operational Documentation Dossier",
        "",
        "Version: 2.1",
        f"Generated: {TODAY}",
        "Prepared for: IMD/CWC Review, Academic Submission, Internship Evaluation, Hackathon Presentation, Production Handover, Technical Review, and Deployment Documentation.",
        "",
        "## Document Control",
        "",
        md_table(["Field", "Value"], [
            ["Document Type", "IMD/CWC operational review and production handover dossier"],
            ["Evidence Basis", "Supplied UI screenshots, frontend source code, backend source code, Supabase schema, generated migration/deployment reports, and historical archive files"],
            ["Source Repository", str(ROOT)],
            ["Screenshot Evidence", f"{len(SCREENSHOTS)} supplied screenshots copied into `docs/screenshots/`"],
            ["Prepared Date", TODAY],
        ]),
        "",
        "## Version History",
        "",
        md_table(["Version", "Date", "Summary"], [
            ["1.0", "Phase 2/3", "Deterministic convective forecasting, archive, verification, and operations modules."],
            ["1.5", "Phase 5.9", "Reviewer polish, file analysis center, registry, typography, and deployment refinement."],
            ["2.0", TODAY, "Supabase migration, single-URL Render deployment, and complete IMD/CWC documentation dossier."],
            ["2.1", TODAY, "Screenshot-backed evidence catalogue, standalone audit deliverables, and front-matter simplified for reviewer use."],
        ]),
        "",
        "## Acknowledgement",
        "StormSense AI was developed as an operational atmospheric intelligence workstation focused on thunderstorm forecasting, verification, historical analysis, and IMD/CWC-style review workflows. The system integrates live/cached sounding ingestion, deterministic forecast engines, historical archives, verification science, and production deployment configuration.",
        "",
        "## Executive Summary",
        f"StormSense AI is a FastAPI + Vite React atmospheric decision-support platform. Repository inspection found {COUNTS['total_files']} source/data/runtime files under `backend/` and `frontend/src/`, {len(ROUTES)} FastAPI route decorators, {len(SCHEMA)} Supabase schema tables, {COUNTS['historical_archive_records']} RSRW historical archive records, and {COUNTS['observational_records']} IMD observational records. The application provides radar-like monitoring, deterministic forecast generation, historical thunderstorm archive workflows, file upload analysis, verification metrics, district impact intelligence, operational bulletins, review mode, and single-URL Render deployment.",
        "",
        "## Table of Contents",
        "1. Project Overview\n2. System Architecture\n3. Page-by-Page Documentation\n4. Screenshot Evidence Catalogue\n5. Research & Insights Documentation\n6. Meteorological Glossary\n7. Database Documentation\n8. API Documentation\n9. Deployment Documentation\n10. Testing Documentation\n11. User Manual\n12. Developer Manual\n13. IMD Demonstration Guide\n14. Appendix: Project Inventory",
        "",
        "# 1. Project Overview",
        "",
        "## Introduction",
        "StormSense AI is an IMD-style atmospheric intelligence platform designed to support severe weather monitoring and thunderstorm forecasting decisions. It combines operational frontend workstations with backend forecast, verification, archive, and research engines.",
        "",
        "## Project Objectives",
        "- Interpret sounding and historical weather datasets in operational terms.\n- Produce deterministic thunderstorm forecasts and probability diagnostics.\n- Verify forecasts scientifically using CSI, POD, FAR, HSS, and BIAS.\n- Support IMD/CWC review through simplified review mode and reviewer dashboard.\n- Preserve historical archive workflows for 2023, 2024, and 2025 records.\n- Serve frontend and backend from one Render URL for production readiness.",
        "",
        "## Problem Statement",
        "Thunderstorm forecasting review requires a joined workflow: sounding interpretation, convective index calculation, forecast reasoning, observed-event verification, district impact assessment, and reviewer decision capture. Spreadsheet-only or display-only systems make it difficult to trace why a forecast was issued, whether it verified, and how thresholds should be tuned.",
        "",
        "## Need for Thunderstorm Forecast Verification",
        "Operational confidence depends on observed-event validation. StormSense AI provides contingency analysis, analog comparison, threshold reliability, archive case review, and probability traceability so that forecast behavior can be reviewed scientifically rather than treated as opaque dashboard output.",
        "",
        "## IMD/CWC Operational Relevance",
        "The system emphasizes Visakhapatnam, Machilipatnam, north coastal Andhra, radiosonde cycle metadata, district-level impact narratives, lightning/heavy rain/squall guidance, and operational bulletin workflows. It uses station codes, synoptic cycles, deterministic thresholds, and IMD-style terms throughout the review experience.",
        "",
        "## Scope of StormSense AI",
        "The platform covers live/cached sounding ingestion, forecast generation, radar-sounding-map fusion, historical archive analysis, file upload analysis, probabilistic forecasting, verification science, climatology, analog intelligence, reviewer dashboards, and single-service deployment.",
        "",
        "## Key Capabilities",
        md_table(["Capability", "Implementation Evidence"], [
            ["Live/cached sounding ingestion", "`backend/fetch_sounding.py`, `/cwc/sounding-raw/{station_code}`, CAPE traceability metadata"],
            ["Forecast engine", "`backend/main.py`, `fetch_station_data`, `/forecast`, `PredictorEngine.jsx`"],
            ["Verification science", "`backend/analysis_engines.py`, `/cwc/verification`, `/cwc/threshold-research`, `ResearchHub.jsx`"],
            ["Historical archive", "`backend/data/rsrw_historical_archive.*`, `/cwc/historical-dates`, `ResearchHub.jsx:HISTORICAL_WORKBENCH`"],
            ["File analysis center", "`/cwc/analyze-historical-dataset`, `ResearchHub.jsx`, `Phase3OpsModule.jsx`"],
            ["District impact", "`compute_district_impact`, `/cwc/district-impact`, Andhra monitoring module"],
            ["Single URL deployment", "`render.yaml`, static serving in `backend/main.py`"],
        ]),
        "",
        "# 2. System Architecture",
        "",
        architecture_diagrams(),
        "",
        "## Frontend Architecture",
        f"Primary navigation is implemented in `Sidebar.jsx` with {len(SIDEBAR)} command deck entries. `App.jsx` manages authentication, review mode, forecast/historical/trend/escalation state, websocket connection, and module routing.",
        "",
        md_table(["Navigation ID", "Visible Label"], sidebar_rows),
        "",
        "## Backend Architecture",
        "The FastAPI backend is centered in `backend/main.py`. It owns authentication routes, forecast routes, websocket stream, system status, CWC/research endpoints, file upload endpoints, and static frontend serving. Scientific logic is delegated to `analysis_engines.py`, `thermo.py`, and `fetch_sounding.py`.",
        "",
        "## Engine Architecture",
        md_table(["Engine", "Source", "Role"], [
            ["Thermodynamic solver", "backend/thermo.py", "Sounding parsing, CAPE/CIN/LCL/LFC/EL and parcel-path calculations."],
            ["Sounding ingestion", "backend/fetch_sounding.py", "Wyoming fetch, cache metadata, cycle-aware source status, freshness scoring."],
            ["Research and verification", "backend/analysis_engines.py", "Historical observations, contingency metrics, threshold research, probability, climatology, analogs, district impact."],
            ["Forecast orchestration", "backend/main.py", "Station forecast assembly, lifecycle, decision support, persistence, websocket output."],
            ["ML training artifact", "backend/ml/train_model.py and storm_model.pkl", "Random forest training script and model artifact for storm classification support."],
        ]),
        "",
        page_docs(),
        "",
        "# 4. Screenshot Evidence Catalogue",
        "",
        screenshot_evidence_section(),
        "",
        "# 5. Research & Insights Documentation",
        "",
        md_table(["Tab Code", "Visible Page", "Operational Function"], research_rows),
        "",
        "The Research & Insights module is the primary review and investigation area. It connects historical archive records, forecast simulation, verification, dataset upload analysis, thunderstorm registry, operational metadata, CAPE diagnostics, probability evolution, and reviewer workflow into a single operational research desk.",
        "",
        "# 6. Meteorological Glossary",
        "",
        md_table(["Term", "Operational Interpretation"], glossary_rows),
        "",
        "# 7. Database Documentation",
        "",
        "StormSense AI uses Supabase PostgreSQL for production persistence. The schema is defined in `supabase_schema.sql` and initialized/accessed through SQLAlchemy helpers.",
        "",
        md_table(["Table", "Columns"], schema_rows),
        "",
        code_block("mermaid", """
erDiagram
  users ||--o{ profiles : owns
  users ||--o{ audit_login : records
  thunderstorm_forecasts ||--o{ verification_results : validates
  historical_observations ||--o{ thunderstorm_registry : classifies
  uploaded_datasets ||--o{ thunderstorm_registry : generates
  reviewer_dashboard_records }o--|| users : reviewed_by
        """),
        "",
        database_details_for_main(),
        "",
        "# 8. API Documentation",
        "",
        f"Full API reference is also generated separately in `API_Reference.md`. Summary of discovered routes ({len(ROUTES)}):",
        "",
        md_table(["Method", "Route", "Handler", "Source"], route_rows),
        "",
        api_details_for_main(),
        "",
        "# 9. Deployment Documentation",
        "",
        "Full deployment guide is generated separately in `Deployment_Guide.md`. The current production direction is a single Render Web Service. `render.yaml` builds backend requirements, builds the Vite frontend, and starts `uvicorn backend.main:app`.",
        "",
        "# 10. Testing Documentation",
        "",
        md_table(["Test Area", "Validation Method", "Expected Result"], [
            ["Frontend build", "`npm install`; `npm run build`", "`frontend/dist` generated successfully."],
            ["Backend compile", "`python -m py_compile backend/main.py`", "No syntax errors."],
            ["Backend import", "`python -c \"import backend.main\"`", "Application imports without ModuleNotFoundError."],
            ["API smoke", "GET `/system-status`, `/history`, `/cwc/historical-dates`, `/forecast`", "HTTP 200 JSON responses."],
            ["SPA refresh", "GET `/dashboard`, `/research`, `/archive`, `/forecast` with `Accept: text/html`", "HTTP 200 React HTML shell."],
            ["Upload workflow", "POST `/cwc/analyze-historical-dataset` with CSV/XLS/XLSX", "Dataset quality summary and registry returned."],
            ["Verification workflow", "Run threshold research and contingency analysis", "CSI, HSS, POD, FAR, BIAS returned."],
            ["Review mode", "Enable IMD Review Mode", "Simplified review deck with only reviewer-critical fields."],
        ]),
        "",
        testing_details_for_main(),
        "",
        "# 11. User Manual",
        "",
        user_manual(),
        "",
        "# 12. Developer Manual",
        "",
        developer_manual(),
        "",
        "# 13. IMD Demonstration Guide",
        "",
        "## Demo Scenario 1: Historical Event Verification",
        "Open Research & Insights -> Historical Thunderstorm Archive, select a 2025 or 2024 event, inspect observed event, forecast reproduction, verification result, and threshold trace. Talking point: archive records participate in latest-event logic and thunderstorm registry filtering.",
        "",
        "## Demo Scenario 2: Forecast Simulation",
        "Open Thunderstorm Forecast Simulator, select station and inputs, run simulation, and explain CAPE/LI/PWAT/SWEAT/K-index influence on forecast classification and recommended action.",
        "",
        "## Demo Scenario 3: Dataset Upload Analysis",
        "Upload a CSV/XLS/XLSX historical dataset through Analyze Historical Dataset. Review records, missing values, duplicates, quality score, stations detected, date range, registry, and verification.",
        "",
        "## Demo Scenario 4: Operational Monitoring",
        "Open Coastal Thunderstorm Monitoring Center or Live Operational Nowcast Center. Demonstrate Visakhapatnam/Machilipatnam focus, marine inflow, convergence, lightning/rainfall corridors, district impacts, and squall propagation guidance.",
        "",
        "## Expected Outputs",
        "- Forecast classification and probability.\n- Meteorologist explanation and recommended action.\n- Verification result and skill metrics.\n- Thunderstorm registry records.\n- Exportable reports and bulletins.",
        "",
        "## Reviewer Questions",
        md_table(["Question", "Suggested Evidence"], [
            ["Where does CAPE come from?", "CAPE traceability endpoint, sounding metadata, cache/source status."],
            ["How are forecasts verified?", "Contingency matrix, CSI/POD/FAR/HSS/BIAS, historical analysis workflow."],
            ["Can uploaded datasets be reviewed?", "File Analysis Center and `/cwc/analyze-historical-dataset`."],
            ["How does deployment work?", "`render.yaml`, single FastAPI service, static SPA fallback."],
            ["Does the system preserve IMD review usability?", "Review Mode, Reviewer Dashboard, Start Here, About page."],
        ]),
        "",
        "# 14. Appendix: Project Inventory",
        "",
        md_table(["Item", "Value"], [
            ["Backend/frontend source inventory", str(COUNTS["extensions"])],
            ["FastAPI route decorators", str(len(ROUTES))],
            ["Supabase tables", str(len(SCHEMA))],
            ["RSRW historical archive records", str(COUNTS["historical_archive_records"])],
            ["IMD observational records", str(COUNTS["observational_records"])],
            ["Checked-in image assets", ", ".join(COUNTS["image_assets"]) or "None"],
        ]),
        "",
    ]
    return "\n".join(text)


def write_markdown_files() -> dict[str, str]:
    files = {
        "MVP_DOCUMENTATION_VERSION.md": mvp_documentation_version(),
        "FULL_DOCUMENTATION_EXPANSION_PLAN.md": full_documentation_expansion_plan(),
        "Architecture_Diagrams.md": architecture_diagrams(),
        "SCREENSHOT_ANALYSIS_REPORT.md": screenshot_analysis_report(),
        "REQUIREMENTS_SPECIFICATION.md": requirements_specification(),
        "DESIGN_SYSTEM_GUIDE.md": design_system_guide(),
        "API_Reference.md": api_reference(),
        "ACCESSIBILITY_AUDIT.md": accessibility_audit(),
        "SECURITY_GUIDE.md": security_guide(),
        "TESTING_REPORT.md": testing_report(),
        "User_Manual.md": user_manual(),
        "Developer_Manual.md": developer_manual(),
        "Deployment_Guide.md": deployment_guide(),
        "StormSense_AI_IMD_Documentation.md": main_document(),
    }
    for name, content in files.items():
        (ROOT / name).write_text(content + "\n", encoding="utf-8")
    return files


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "DADCE0")
        borders.append(tag)
    tbl_pr.append(borders)


def add_markdown_table(doc: Document, lines: list[str]):
    rows = []
    for line in lines:
        if re.match(r"^\|\s*-", line):
            continue
        cells = [c.strip().replace("\\|", "|") for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=cols)
    table.autofit = True
    set_table_borders(table)
    for r_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for c_idx in range(cols):
            cells[c_idx].text = row[c_idx] if c_idx < len(row) else ""
            for para in cells[c_idx].paragraphs:
                para.paragraph_format.space_after = Pt(2)
                for run in para.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(8.5)
            if r_idx == 0:
                set_cell_shading(cells[c_idx], "E8EEF5")
                for para in cells[c_idx].paragraphs:
                    for run in para.runs:
                        run.bold = True
    doc.add_paragraph("")


def markdown_to_docx(markdown_text: str, out_path: Path):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)
    for style_name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)

    footer = section.footer.paragraphs[0]
    footer.text = "StormSense AI IMD/CWC Operational Documentation Dossier | Version 2.1"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    lines = markdown_text.splitlines()
    i = 0
    in_code = False
    code_lines = []
    table_lines = []
    first_h1 = True
    while i < len(lines):
        line = lines[i]
        if line.startswith("```"):
            if in_code:
                para = doc.add_paragraph()
                run = para.add_run("\n".join(code_lines))
                run.font.name = "Courier New"
                run.font.size = Pt(8)
                para.paragraph_format.left_indent = Inches(0.15)
                para.paragraph_format.space_after = Pt(6)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue
        if line.startswith("|"):
            table_lines.append(line)
            i += 1
            if i >= len(lines) or not lines[i].startswith("|"):
                add_markdown_table(doc, table_lines)
                table_lines = []
            continue
        if not line.strip():
            i += 1
            continue
        img_match = re.match(r"^!\[(.*?)\]\((.*?)\)", line)
        if img_match:
            caption, rel_path = img_match.group(1), img_match.group(2)
            image_path = (ROOT / rel_path).resolve()
            if image_path.exists():
                try:
                    doc.add_picture(str(image_path), width=Inches(6.2))
                    cap = doc.add_paragraph(caption)
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in cap.runs:
                        run.italic = True
                        run.font.size = Pt(8)
                except Exception as exc:
                    doc.add_paragraph(f"[Image unavailable: {rel_path} ({exc})]")
            else:
                doc.add_paragraph(f"[Image not found: {rel_path}]")
            i += 1
            continue
        if line.startswith("# "):
            if not first_h1:
                doc.add_page_break()
            first_h1 = False
            para = doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue
        if line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=3)
            i += 1
            continue
        if re.match(r"^\d+\.\s+", line):
            doc.add_paragraph(re.sub(r"^\d+\.\s+", "", line), style="List Number")
            i += 1
            continue
        if line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
            i += 1
            continue
        para = doc.add_paragraph(line)
        para.paragraph_format.line_spacing = 1.25
        para.paragraph_format.space_after = Pt(6)
        i += 1

    doc.save(out_path)


def markdown_to_pdf(markdown_text: str, out_path: Path):
    styles = getSampleStyleSheet()
    body = ParagraphStyle("BodyCompact", parent=styles["BodyText"], fontName="Helvetica", fontSize=9.2, leading=12, spaceAfter=5)
    h1 = ParagraphStyle("H1", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=15, leading=18, textColor=colors.HexColor("#2E74B5"), spaceAfter=8)
    h2 = ParagraphStyle("H2", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=12.5, leading=15, textColor=colors.HexColor("#2E74B5"), spaceAfter=6)
    h3 = ParagraphStyle("H3", parent=styles["Heading3"], fontName="Helvetica-Bold", fontSize=11, leading=13, textColor=colors.HexColor("#1F4D78"), spaceAfter=5)
    code = ParagraphStyle("Code", parent=styles["Code"], fontName="Courier", fontSize=7.5, leading=9, spaceAfter=6)
    story = []
    lines = markdown_text.splitlines()
    i = 0
    in_code = False
    code_lines = []
    table_lines = []
    first_h1 = True
    while i < len(lines):
        line = lines[i]
        if line.startswith("```"):
            if in_code:
                story.append(Preformatted("\n".join(code_lines), code))
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue
        if line.startswith("|"):
            table_lines.append(line)
            i += 1
            if i >= len(lines) or not lines[i].startswith("|"):
                rows = []
                for tline in table_lines:
                    if re.match(r"^\|\s*-", tline):
                        continue
                    rows.append([c.strip() for c in tline.strip().strip("|").split("|")])
                if rows:
                    table = Table(rows, repeatRows=1)
                    table.setStyle(TableStyle([
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF5")),
                        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#DADCE0")),
                        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                        ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
                        ("FONTSIZE", (0, 0), (-1, -1), 6.2),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 3),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 3),
                        ("TOPPADDING", (0, 0), (-1, -1), 3),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                    ]))
                    story.append(table)
                    story.append(Spacer(1, 0.08 * inch))
                table_lines = []
            continue
        if not line.strip():
            i += 1
            continue
        img_match = re.match(r"^!\[(.*?)\]\((.*?)\)", line)
        if img_match:
            caption, rel_path = img_match.group(1), img_match.group(2)
            image_path = (ROOT / rel_path).resolve()
            if image_path.exists():
                try:
                    image = RLImage(str(image_path))
                    max_width = 6.3 * inch
                    scale = min(1.0, max_width / float(image.imageWidth))
                    image.drawWidth = image.imageWidth * scale
                    image.drawHeight = image.imageHeight * scale
                    story.append(image)
                    story.append(Paragraph(caption.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"), body))
                    story.append(Spacer(1, 0.12 * inch))
                except Exception as exc:
                    story.append(Paragraph(f"[Image unavailable: {rel_path} ({exc})", body))
            else:
                story.append(Paragraph(f"[Image not found: {rel_path}]", body))
            i += 1
            continue
        safe = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        if line.startswith("# "):
            if not first_h1:
                story.append(PageBreak())
            first_h1 = False
            story.append(Paragraph(safe[2:].strip(), h1))
        elif line.startswith("## "):
            story.append(Paragraph(safe[3:].strip(), h2))
        elif line.startswith("### "):
            story.append(Paragraph(safe[4:].strip(), h3))
        elif line.startswith("#### "):
            story.append(Paragraph(safe[5:].strip(), h3))
        elif line.startswith("- "):
            story.append(Paragraph("• " + safe[2:].strip(), body))
        elif re.match(r"^\d+\.\s+", line):
            story.append(Paragraph(safe, body))
        else:
            story.append(Paragraph(safe, body))
        i += 1

    doc = SimpleDocTemplate(str(out_path), pagesize=letter, leftMargin=0.75 * inch, rightMargin=0.75 * inch, topMargin=0.75 * inch, bottomMargin=0.75 * inch)
    doc.build(story)


def main():
    files = write_markdown_files()
    main_md = files["StormSense_AI_IMD_Documentation.md"]
    docx_path = ROOT / "StormSense_AI_IMD_Documentation.docx"
    pdf_path = ROOT / "StormSense_AI_IMD_Documentation.pdf"
    try:
        markdown_to_docx(main_md, docx_path)
    except PermissionError:
        docx_path = ROOT / "StormSense_AI_IMD_Documentation_v2_1.docx"
        markdown_to_docx(main_md, docx_path)
    try:
        markdown_to_pdf(main_md, pdf_path)
    except PermissionError:
        pdf_path = ROOT / "StormSense_AI_IMD_Documentation_v2_1.pdf"
        markdown_to_pdf(main_md, pdf_path)
    manifest = {
        "generated": TODAY,
        "routes": len(ROUTES),
        "schema_tables": len(SCHEMA),
        "historical_archive_records": COUNTS["historical_archive_records"],
        "observational_records": COUNTS["observational_records"],
        "screenshots": len(SCREENSHOTS),
        "files": list(files.keys()) + [docx_path.name, pdf_path.name],
    }
    (ROOT / "scratch" / "imd_docs_manifest.json").write_text(json.dumps(manifest, indent=2), encoding="utf-8")
    print(json.dumps(manifest, indent=2))


if __name__ == "__main__":
    main()
